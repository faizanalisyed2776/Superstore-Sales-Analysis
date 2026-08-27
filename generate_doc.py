import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def build_assignment_doc():
    doc = docx.Document()

    # --- Page Setup ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Title & Cover Page Info ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FOUNDATIONAL DATA ANALYTICS REPORT:\nEXPLORATION & STRATEGY OUTLINE")
    run.font.name = 'Calibri'
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run("Dataset Exploration, Problem Definition, and Analysis Strategy")
    run_sub.font.name = 'Calibri'
    run_sub.font.size = Pt(14)
    run_sub.font.italic = True
    run_sub.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Helper function for headings
    def add_custom_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        for r in h.runs:
            r.font.name = 'Calibri'
            r.font.color.rgb = RGBColor(0, 51, 102)
        return h

    # --- Section 1: Executive Summary & Overview ---
    add_custom_heading("1. Executive Summary & Dataset Overview")
    
    p = doc.add_paragraph(
        "This foundational report outlines the initial exploratory data analysis (EDA), problem statement, "
        "and strategic roadmap for analyzing the Superstore Retail Sales Dataset (sourced from Kaggle). "
        "The primary goal is to examine transactional records to uncover profitability drivers, evaluate regional "
        "performance disparities, and optimize discount structures."
    )
    p.paragraph_format.space_after = Pt(10)

    p_ds = doc.add_paragraph()
    p_ds.add_run("Dataset Summary:\n").bold = True
    p_ds.add_run("• Source: Public Superstore Sales Dataset (Kaggle)\n")
    p_ds.add_run("• Dimensions: 9,994 records (rows) across 21 attributes (columns)\n")
    p_ds.add_run("• Primary Entities: Customer Segments, Product Categories, Regional Hubs, Order Metrics")

    # --- Section 2: Variable Classification Table ---
    add_custom_heading("2. Variable Classification & Data Types")

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Variable Name"
    hdr_cells[1].text = "Data Type & Measurement Scale"
    hdr_cells[2].text = "Description / Strategic Relevance"

    # Formatting Table Header
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        # Setting background color via XML (optional, plain text fallback below)

    variables_data = [
        ("Ship Mode", "Categorical (Nominal)", "Shipping method (Standard, Second, First, Same Day)"),
        ("Segment", "Categorical (Nominal)", "Target market segment (Consumer, Corporate, Home Office)"),
        ("Region", "Categorical (Nominal)", "Geographic distribution hub (East, West, Central, South)"),
        ("Category / Sub-Category", "Categorical (Nominal)", "Product hierarchy (Furniture, Office Supplies, Technology)"),
        ("Sales", "Numerical (Continuous)", "Gross monetary revenue generated per order line item"),
        ("Quantity", "Numerical (Discrete)", "Number of units purchased per transaction"),
        ("Discount", "Numerical (Continuous)", "Percentage price reduction applied (0.0 to 0.8)"),
        ("Profit", "Numerical (Continuous)", "Net profit or loss realized after costs")
    ]

    for var, dtype, desc in variables_data:
        row_cells = table.add_row().cells
        row_cells[0].text = var
        row_cells[1].text = dtype
        row_cells[2].text = desc

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- Section 3: Exploratory Data Analysis (EDA) Workflow ---
    add_custom_heading("3. Exploratory Data Analysis Workflow & Python Implementation")
    
    doc.add_paragraph(
        "Data hygiene and descriptive statistics were computed using Python (Pandas, Numpy, Seaborn). "
        "The following snippet demonstrates data loading, deduplication, missing value imputation, "
        "and correlation evaluation:"
    )

    code_snippet = (
        "import pandas as pd\n"
        "import numpy as np\n"
        "import seaborn as sns\n"
        "import matplotlib.pyplot as plt\n\n"
        "# 1. Load & Inspect Dataset\n"
        "df = pd.read_csv('superstore_sales.csv')\n"
        "print('Dataset Shape:', df.shape)\n\n"
        "# 2. Cleaning & Missing Value Handling\n"
        "df = df.drop_duplicates()\n"
        "df['Postal Code'] = df['Postal Code'].fillna(df['Postal Code'].mode()[0])\n\n"
        "# 3. Statistical Summary\n"
        "summary_stats = df[['Sales', 'Quantity', 'Discount', 'Profit']].describe()\n\n"
        "# 4. Heatmap & Correlation Matrix\n"
        "plt.figure(figsize=(8, 5))\n"
        "sns.heatmap(df[['Sales', 'Quantity', 'Discount', 'Profit']].corr(), annot=True, cmap='Blues')\n"
        "plt.title('Feature Correlation Matrix')\n"
        "plt.savefig('correlation_matrix.png')"
    )

    code_p = doc.add_paragraph()
    code_run = code_p.add_run(code_snippet)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(9.5)
    code_p.paragraph_format.left_indent = Inches(0.2)

    # --- Section 4: Problem Statement & Hypotheses ---
    add_custom_heading("4. Problem Definition Statement & Analytical Hypotheses")

    doc.add_paragraph(
        "Problem Statement:\n"
        "Despite achieving strong top-line sales growth across regional markets, retail entities experience severe profit "
        "margin erosion due to unmonitored discounting strategies and higher fulfillment overheads in specific product sub-categories."
    ).runs[0].bold = True

    p_hyp = doc.add_paragraph()
    p_hyp.add_run("Key Analytical Questions & Hypotheses:\n").bold = True
    p_hyp.add_run("1. Q1: Which product sub-categories consistently drive negative profit margins despite high sales volume?\n")
    p_hyp.add_run("2. Q2: What is the optimal discount threshold before gross profit margins become negative?\n\n")
    p_hyp.add_run("• Hypothesis 1 (H₁): Order discount rates exceeding 20% cause a statistically significant drop in net profit per order item.\n")
    p_hyp.add_run("• Hypothesis 2 (H₂): Regional profitability varies significantly based on logistics and choice of shipping mode.")

    # --- Section 5: Proposed Analysis Strategy ---
    add_custom_heading("5. Proposed Analysis Strategy & Methodology")

    strat_table = doc.add_table(rows=1, cols=3)
    strat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    strat_table.style = 'Table Grid'

    s_hdr = strat_table.rows[0].cells
    s_hdr[0].text = "Phase"
    s_hdr[1].text = "Proposed Method / Tool"
    s_hdr[2].text = "Expected Milestone / Outcome"

    strategy_data = [
        ("Phase 1: Data Preprocessing", "Pandas (IQR outlier clipping, One-Hot Encoding)", "Clean, transformed dataset with standardized numeric scale."),
        ("Phase 2: Exploratory Visualization", "Seaborn / Matplotlib (Boxplots, FacetGrid)", "Identification of extreme loss-making categories and regional trends."),
        ("Phase 3: Statistical Hypothesis Testing", "SciPy (Two-Sample t-Tests, ANOVA)", "Statistical validation of discounting impact on gross profit margins."),
        ("Phase 4: Predictive Modeling Prep", "Scikit-Learn (Linear & Ridge Regression)", "Baseline regression pipeline for profit and sales forecasting.")
    ]

    for p_name, p_method, p_out in strategy_data:
        r_cells = strat_table.add_row().cells
        r_cells[0].text = p_name
        r_cells[1].text = p_method
        r_cells[2].text = p_out

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # --- Section 6: Limitations & Future Scope ---
    add_custom_heading("6. Data Limitations & Future Work")

    p_lim = doc.add_paragraph()
    p_lim.add_run("Limitations:\n").bold = True
    p_lim.add_run("• The dataset lacks individual customer demographic parameters (e.g., age, household income), restricting detailed persona analysis.\n")
    p_lim.add_run("• External macroeconomic indicators (inflation rates, regional tax dynamics) are omitted, which limits market force modeling.\n\n")
    p_lim.add_run("Future Scope:\n").bold = True
    p_lim.add_run("• Incorporating time-series forecasting (ARIMA / Prophet) to predict seasonal sales surges and optimize inventory replenishment cycles.")

    # Save Document
    doc.save("Data_Analytics_Assignment.docx")
    print("SUCCESS: 'Data_Analytics_Assignment.docx' has been generated and saved to your directory.")

if __name__ == "__main__":
    build_assignment_doc()